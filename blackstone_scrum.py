"""
Blackstone Daily Scrum Report Generator
Fetches data from:
  - Azure DevOps (work items tagged Blackstone)
  - Pendo (visitor engagement for Blackstone segment)
  - Platform metrics (pluggable — fill in your source)

Usage:
  python3 blackstone_scrum.py

Required env vars (or edit the CONFIG block below):
  ADO_PAT          - Azure DevOps Personal Access Token (Work Items: Read)
  PENDO_API_KEY    - Pendo Integration key (Settings > Integrations > Integration keys)
  PENDO_SEGMENT_ID_EU      - Pendo segment ID for Blackstone EU
  PENDO_SEGMENT_ID_USEAST  - Pendo segment ID for Blackstone USEast
"""

import os
import sys
import base64
import json
import datetime
import io
import requests
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.gridspec import GridSpec


def fmt_day(dt) -> str:
    """Format day without leading zero — works on Linux, macOS, Windows."""
    return str(dt.day)


def fmt_date(dt) -> str:
    """e.g. Wednesday, June 18, 2026 — cross-platform."""
    return f"{dt.strftime('%A, %B')} {fmt_day(dt)}, {dt.year}"


def fmt_mon(dt) -> str:
    """e.g. '18 Jun' — cross-platform."""
    return f"{fmt_day(dt)} {dt.strftime('%b')}"


def fmt_month_day(ms) -> str:
    """e.g. 'June 18' (no leading zero) — matches working script output."""
    if not ms:
        return "—"
    d = datetime.datetime.fromtimestamp(ms / 1000, PACIFIC)
    return d.strftime("%B %d").replace(" 0", " ")

# ─── CONFIG ──────────────────────────────────────────────────────────────────

ADO_ORG     = "CoreStack-Tech"
ADO_PROJECT = "Product_Mgmt"
ADO_PAT       = os.environ.get("ADO_PAT", "")
PENDO_API_KEY = os.environ.get("PENDO_API_KEY", "")

# ── Local credentials override (never committed) ──────────────────────────
# Create a file called  credentials.py  in the same folder with:
#   ADO_PAT       = "your-ado-pat"
#   PENDO_API_KEY = "your-pendo-key"
try:
    from credentials import ADO_PAT as _A, PENDO_API_KEY as _P  # type: ignore
    if _A: ADO_PAT = _A
    if _P: PENDO_API_KEY = _P
except ImportError:
    pass
PENDO_REGION_EU      = "portal.corestack.io"
PENDO_REGION_USEAST  = "useast.corestack.io"
PENDO_SEGMENT_ID_EU  = "Hh8nJFk5pC2QUMWBQavVPz3Y9zw"   # Blackstone EU segment
PENDO_EXCLUDED_EMAIL = None   # set to an email string to exclude, or None to include all
PENDO_WINDOW_DAYS    = 3

# ─── ADO ─────────────────────────────────────────────────────────────────────

def ado_headers():
    token = base64.b64encode(f":{ADO_PAT}".encode()).decode()
    return {
        "Authorization": f"Basic {token}",
        "Content-Type": "application/json",
    }

def fetch_ado_blackstone_incidents():
    """
    Runs a WIQL query to fetch all active Blackstone work items.
    Filters by tag = 'Blackstone' and state != Closed/Removed.
    Returns list of dicts with id, title, assignedTo, state, priority, tags, workItemType.
    """
    if not ADO_PAT:
        # Return fake data so the script runs without a real PAT
        return _fake_ado_data()

    url = f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}/_apis/wit/wiql?api-version=7.1"

    wiql = {
        "query": """
            SELECT [System.Id], [System.Title], [System.AssignedTo],
                   [System.State], [Microsoft.VSTS.Common.Priority],
                   [System.Tags], [System.WorkItemType],
                   [System.AreaPath]
            FROM WorkItems
            WHERE [System.TeamProject] = @project
              AND [System.Tags] CONTAINS 'Blackstone'
              AND [System.WorkItemType] = 'Incident'
              AND [System.State] IN ('New', 'In Progress', 'Awaiting Deployment')
            ORDER BY [Microsoft.VSTS.Common.Priority] ASC,
                     [System.ChangedDate] DESC
        """
    }

    resp = requests.post(url, headers=ado_headers(), json=wiql)
    resp.raise_for_status()
    ids = [str(item["id"]) for item in resp.json().get("workItems", [])]

    if not ids:
        return []

    # Batch fetch details (max 200 per call)
    details_url = (
        f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}/_apis/wit/workitems"
        f"?ids={','.join(ids[:200])}"
        f"&fields=System.Id,System.Title,System.AssignedTo,System.State,"
        f"Microsoft.VSTS.Common.Priority,System.Tags,System.WorkItemType,System.AreaPath"
        f"&api-version=7.1"
    )
    resp2 = requests.get(details_url, headers=ado_headers())
    resp2.raise_for_status()

    results = []
    for wi in resp2.json().get("value", []):
        f = wi["fields"]
        assigned = f.get("System.AssignedTo", {})
        results.append({
            "id":           wi["id"],
            "title":        f.get("System.Title", ""),
            "assignedTo":   assigned.get("displayName", "Unassigned") if isinstance(assigned, dict) else str(assigned),
            "state":        f.get("System.State", ""),
            "priority":     f.get("Microsoft.VSTS.Common.Priority", ""),
            "tags":         f.get("System.Tags", ""),
            "workItemType": f.get("System.WorkItemType", ""),
            "areaPath":     f.get("System.AreaPath", ""),
        })
    return results

def _fake_ado_data():
    """Placeholder data — remove once ADO_PAT is set."""
    return [
        {
            "id": 134003,
            "title": "Assessment Report UI Error Despite Successful Report Generation",
            "assignedTo": "Zubair",
            "state": "In Progress",
            "priority": 2,
            "tags": "Blackstone; Analytics",
            "workItemType": "Bug",
            "areaPath": "Product_Mgmt\\Analytics",
        },
        {
            "id": 131077,
            "title": "Blackstone - Rover - Forecast data missing",
            "assignedTo": "Aadhithya Shanmugapriyan",
            "state": "New",
            "priority": 2,
            "tags": "Blackstone; FinOps",
            "workItemType": "Bug",
            "areaPath": "Product_Mgmt\\FinOps",
        },
        {
            "id": 131076,
            "title": "Blackstone - RBAC permission issue on tenant view",
            "assignedTo": "Aadhithya Shanmugapriyan",
            "state": "Awaiting Deployment",
            "priority": 3,
            "tags": "Blackstone; Core",
            "workItemType": "Bug",
            "areaPath": "Product_Mgmt\\Core",
        },
    ]

def bundle_from_area(area_path: str) -> str:
    mapping = {
        "FinOps": "FinOps",
        "Analytics": "Analytics",
        "CloudOps": "CloudOps",
        "Core": "Core",
    }
    for key, val in mapping.items():
        if key.lower() in area_path.lower():
            return val
    return "—"

# ─── PENDO ───────────────────────────────────────────────────────────────────

PENDO_ENV_VALUES  = ["Blackstone-USeast", "Blackstone-Useast"]
PACIFIC           = datetime.timezone(datetime.timedelta(hours=-7))  # PDT

def pendo_headers():
    return {
        "x-pendo-integration-key": PENDO_API_KEY,
        "Content-Type": "application/json",
    }


def _pendo_agg(pipeline, label="", rows_per_page=5000):
    url = "https://app.pendo.io/api/v1/aggregation"
    all_results = []
    start_row = 0
    while True:
        payload = {
            "response": {"mimeType": "application/json", "rowsPerPage": rows_per_page, "startRow": start_row},
            "request":  {"pipeline": pipeline},
        }
        resp = requests.post(url, headers=pendo_headers(), json=payload)
        if not resp.ok:
            print(f"  [Pendo{' ' + label if label else ''}] {resp.status_code}: {resp.text[:400]}")
            break
        data = resp.json()
        page_results = data.get("results", [])
        all_results.extend(page_results)
        total = data.get("total", len(all_results))
        if len(all_results) >= total or not page_results:
            break
        start_row += rows_per_page
    return all_results


def _fetch_blackstone_accounts() -> dict:
    """
    Resolve the 2 Blackstone accounts via metadata.custom.environment.
    Returns {accountId: {"region": ..., "name": ...}}.
      Blackstone-USeast  → Private Equity (USEast)
      Blackstone-Useast  → Real Estate (EU/portal)
    """
    flt = " || ".join(f'metadata.custom.environment=="{v}"' for v in PENDO_ENV_VALUES)
    rows = _pendo_agg([
        {"source": {"accounts": None}},
        {"filter": f"({flt})"},
        {"select": {"accountId": "accountId", "env": "metadata.custom.environment"}},
    ], "accounts")
    result = {}
    for r in rows:
        aid = r.get("accountId")
        env = r.get("env") or ""
        # USeast (capital E) = Private Equity; Useast (lower e) = Real Estate
        if env.endswith("USeast"):
            result[aid] = {"region": "useast", "name": "Private Equity"}
        else:
            result[aid] = {"region": "eu",     "name": "Real Estate"}
    print(f"  [Pendo] Blackstone accounts found: {[(aid, v['name']) for aid, v in result.items()]}")
    return result


def _fetch_eu_segment_vids() -> set:
    """Return the set of visitorIds that belong to the Blackstone EU segment."""
    if not PENDO_SEGMENT_ID_EU:
        return set()
    rows = _pendo_agg([
        {"source": {"visitors": {"segmentId": PENDO_SEGMENT_ID_EU}}},
        {"select": {"visitorId": "visitorId"}},
    ], "eu_segment")
    vids = {r["visitorId"] for r in rows if r.get("visitorId")}
    print(f"  [Pendo] EU segment members: {len(vids)}")
    return vids


def _discover_apps(acct_ids: list) -> list:
    """
    Find every Pendo app touched by the segment by inspecting auto_<appId>
    metadata keys on visitors — matches the dashboard's All-Apps numbers.
    """
    import re
    apps = set()
    for aid in acct_ids:
        rows = _pendo_agg([
            {"source": {"visitors": None}},
            {"filter": f'metadata.auto.accountid=="{aid}"'},
        ], f"apps_{aid[:8]}", rows_per_page=50)
        for r in rows:
            for key in (r.get("metadata") or {}):
                m = re.fullmatch(r"auto_(_?\d+)", key)
                if not m:
                    continue
                raw = m.group(1)
                apps.add(-int(raw[1:]) if raw.startswith("_") else int(raw))
    return sorted(apps) if apps else [None]


def fetch_pendo_all_visitors(accounts: dict = None, window_days: int = PENDO_WINDOW_DAYS):
    """
    Fetch all Blackstone segment members + their activity for the window.
    Logic mirrors generate_blackstone_pendo_report.py exactly:
      - members listed per-account (same visitor appears once per Blackstone account)
      - events/minutes/days summed across all discovered apps (hourRange, matches dashboard)
      - display name: full name if it has a space, else email prefix
      - last seen: "June 18" format (US Pacific)
      - support email included (PENDO_EXCLUDED_EMAIL=None by default)
    """
    if not PENDO_API_KEY:
        return _fake_pendo_visitors()

    accounts  = accounts or {}
    acct_set  = set(accounts.keys())
    if not acct_set:
        return []

    # ── EU segment: authoritative list of EU visitor IDs ─────────────────────
    eu_vids = _fetch_eu_segment_vids()

    # ── Window: US Pacific day buckets, same as dashboard ────────────────────
    now      = datetime.datetime.now(datetime.timezone.utc)
    today_pt = datetime.datetime.now(PACIFIC).date()
    start_pt = datetime.datetime.combine(
        today_pt - datetime.timedelta(days=window_days - 1),
        datetime.time(0, 0), PACIFIC)
    start_ms = int(start_pt.timestamp() * 1000)
    end_ms   = int(now.timestamp() * 1000)

    # ── Step 1: list all segment members ─────────────────────────────────────
    # visitor source uses metadata.auto.accountid; events source uses top-level accountId
    visitor_filter = " || ".join(f'metadata.auto.accountid=="{a}"' for a in acct_set)
    events_filter  = " || ".join(f'accountId=="{a}"' for a in acct_set)
    member_rows = _pendo_agg([
        {"source": {"visitors": None}},
        {"filter": visitor_filter},
        {"select": {
            "visitorId": "visitorId",
            "email":     "metadata.agent.email",
            "name":      "metadata.agent.name",
            "last":      "metadata.auto.lastvisit",
            "accountId": "metadata.auto.accountid",
            "server":    "metadata.auto.lastservername",
        }},
    ], "members")
    print(f"  [Pendo] segment members found: {len(member_rows)}")

    member_vids = {r["visitorId"] for r in member_rows if r.get("visitorId")}

    # ── Step 2: sum events/minutes/days across all apps ───────────────────────
    apps = _discover_apps(list(acct_set))
    print(f"  [Pendo] apps (All Apps): {apps}")
    from collections import defaultdict
    ev_sum   = defaultdict(int)
    min_sum  = defaultdict(int)
    day_sets = defaultdict(set)
    ts = {"period": "hourRange", "first": start_ms, "last": end_ms}
    for app in apps:
        src = {"events": ({"appId": app} if app is not None else None), "timeSeries": ts}
        rows = _pendo_agg([
            {"source": src},
            {"filter": events_filter},   # top-level accountId for events source
            {"identified": "visitorId"},
        ], f"events_app_{app}")
        for r in rows:
            vid = r.get("visitorId")
            if vid not in member_vids:
                continue
            ev_sum[vid]  += r.get("numEvents", 0)
            min_sum[vid] += r.get("numMinutes", 0)
            hour_ms = r.get("hour")
            if hour_ms:
                day_sets[vid].add(
                    datetime.datetime.fromtimestamp(hour_ms / 1000, PACIFIC).strftime("%Y-%m-%d"))
    days_active = {v: len(d) for v, d in day_sets.items()}
    print(f"  [Pendo] visitors active in window: {len(ev_sum)}")
    servers = {(r.get("server") or "").lower() for r in member_rows}
    print(f"  [Pendo] distinct lastservername values: {sorted(servers)[:10]}")
    # Sample: show first 5 visitors with their server and EU-segment membership
    for r in member_rows[:5]:
        vid_s  = r.get("visitorId", "")[:20]
        email_s = (r.get("email") or r.get("visitorId") or "")[:30]
        srv_s  = (r.get("server") or "—")
        in_eu  = vid_s in eu_vids or r.get("visitorId") in eu_vids
        print(f"    sample → {email_s:<30}  server={srv_s:<30}  in_eu_seg={in_eu}")

    # ── Step 3: build result rows (one per member row = one per account) ──────
    results = []
    for meta in member_rows:
        vid   = meta.get("visitorId")
        email = (meta.get("email") or "").strip()
        if not vid:
            continue
        if PENDO_EXCLUDED_EMAIL and email == PENDO_EXCLUDED_EMAIL:
            continue

        # mirrors working script's display_name(): full name if multi-word, else email prefix
        name = (meta.get("name") or "").strip()
        if name and " " in name:
            display = name
        elif email and "@" in email:
            display = email.split("@")[0]
        else:
            display = name or vid
        domain = email.split("@")[1] if "@" in email else "—"

        aid          = meta.get("accountId") or ""
        acct         = accounts.get(aid, {})
        account_name = acct.get("name", "")
        # Account map is most reliable (Private Equity=useast, Real Estate=eu)
        # Fall back to lastservername, then EU segment membership
        server = (meta.get("server") or "").lower()
        acct_region = acct.get("region", "")
        if acct_region:
            region = acct_region
        elif PENDO_REGION_EU in server:
            region = "eu"
        elif PENDO_REGION_USEAST in server:
            region = "useast"
        elif vid in eu_vids:
            region = "eu"
        else:
            region = "unknown"

        last_seen  = fmt_month_day(meta.get("last"))   # "June 18" format
        num_events = ev_sum.get(vid, 0)
        results.append({
            "visitorId":  vid,
            "visitor":    display,
            "domain":     domain,
            "account":    account_name,
            "events":     num_events if num_events else "-",
            "daysActive": days_active.get(vid, 0) if num_events else "-",
            "minutes":    int(min_sum.get(vid, 0)) if num_events else "-",
            "lastSeen":   last_seen,
            "region":     region,
        })

    # active first (events desc), then inactive alphabetically
    results.sort(key=lambda r: (-(r["events"] if isinstance(r["events"], int) else 0),
                                str(r["visitor"]).lower()))
    print(f"  [Pendo] total Blackstone visitor rows: {len(results)}")
    return results


def split_visitors_by_region(visitors):
    eu     = [v for v in visitors if v["region"] == "eu"]
    useast = [v for v in visitors if v["region"] == "useast"]
    other  = [v for v in visitors if v["region"] not in ("eu", "useast")]
    return eu, useast, other


def fetch_pendo_top_pages(accounts: dict = None, window_days: int = PENDO_WINDOW_DAYS, top_n: int = 10):
    """Top pages for Blackstone accounts using pageEvents across all apps."""
    if not PENDO_API_KEY:
        return _fake_pendo_pages()

    accounts  = accounts or {}
    acct_set  = set(accounts.keys())
    if not acct_set:
        return []

    now      = datetime.datetime.now(datetime.timezone.utc)
    today_pt = datetime.datetime.now(PACIFIC).date()
    start_pt = datetime.datetime.combine(
        today_pt - datetime.timedelta(days=window_days - 1), datetime.time(0, 0), PACIFIC)
    start_ms = int(start_pt.timestamp() * 1000)
    end_ms   = int(now.timestamp() * 1000)

    acct_filter = " || ".join(f'accountId=="{a}"' for a in acct_set)
    apps        = _discover_apps(list(acct_set))
    ts          = {"period": "hourRange", "first": start_ms, "last": end_ms}

    page_views: dict = {}
    for app in apps:
        src = {"pageEvents": ({"appId": app} if app is not None else None), "timeSeries": ts}
        rows = _pendo_agg([
            {"source": src},
            {"filter": acct_filter},
            {"group": {"group": ["pageId"], "fields": [{"views": {"sum": "numEvents"}}]}},
        ], f"pages_app_{app}")
        for r in rows:
            pid = r.get("pageId")
            if not pid or pid in ("allevents", "allfeatures"):
                continue
            page_views[pid] = page_views.get(pid, 0) + (r.get("views") or 0)

    sorted_pages = sorted(page_views.items(), key=lambda x: x[1], reverse=True)[:top_n]
    if not sorted_pages:
        return []

    # Step 1: bulk fetch pages for default app + each discovered app
    all_pages = {}
    endpoints = ["https://app.pendo.io/api/v1/page"]
    for app in apps:
        if app is not None:
            endpoints.append(f"https://app.pendo.io/api/v1/app/{app}/page")
    for ep in endpoints:
        try:
            resp = requests.get(ep, headers=pendo_headers(), timeout=15)
            if resp.ok:
                data = resp.json()
                if isinstance(data, list):
                    for p in data:
                        pid = p.get("id")
                        name = p.get("name") or ""
                        if pid and name:
                            all_pages[pid] = name
        except Exception:
            pass

    # Step 2: for any still-unresolved IDs, fetch individually
    unresolved = [pid for pid, _ in sorted_pages if pid not in all_pages]
    for pid in unresolved:
        try:
            resp = requests.get(
                f"https://app.pendo.io/api/v1/page/{requests.utils.quote(str(pid), safe='')}",
                headers=pendo_headers(), timeout=10)
            if resp.ok:
                name = resp.json().get("name") or ""
                if name:
                    all_pages[pid] = name
        except Exception:
            pass

    result = []
    for pid, v in sorted_pages:
        name = all_pages.get(pid, "")
        if name:                          # skip pages with no resolvable name
            result.append({"page": name, "views": v})
    return result[:top_n]


def _fake_pendo_visitors():
    return [
        {"visitor": "robert young",          "domain": "bluemantis.com",  "events": 152, "daysActive": 1, "minutes": 10, "lastSeen": "16 Jun", "region": "useast"},
        {"visitor": "adam schutska",         "domain": "bluemantis.com",  "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "15 Jun", "region": "useast"},
        {"visitor": "bhavana prabhuswamy",   "domain": "bluemantis.com",  "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "15 Jun", "region": "useast"},
        {"visitor": "rj gravel",             "domain": "bluemantis.com",  "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "19 May", "region": "useast"},
        {"visitor": "abagchi",               "domain": "corestack.io",    "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "2 Jun",  "region": "useast"},
        {"visitor": "cspbillingapi",         "domain": "bluemantis.com",  "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "26 May", "region": "useast"},
        {"visitor": "timo pantsari",         "domain": "blackstone.com",  "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "15 Jun", "region": "eu"},
        {"visitor": "alex",                  "domain": "aliando.com",     "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "1 Jun",  "region": "eu"},
        {"visitor": "dipali koche",          "domain": "bluemantis.com",  "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "1 Jun",  "region": "eu"},
        {"visitor": "chris",                 "domain": "aliando.com",     "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "4 Jun",  "region": "eu"},
        {"visitor": "dene donovan",          "domain": "ingrammicro.com", "events": "—", "daysActive": "—", "minutes": "—", "lastSeen": "9 Jun",  "region": "eu"},
    ]

def _fake_pendo_pages():
    return [
        {"page": "Cloud Account Governance", "views": 52},
        {"page": "Unified Governance",       "views": 9},
        {"page": "Inventory",                "views": 4},
        {"page": "Tenants",                  "views": 2},
        {"page": "Optimize Rate",            "views": 2},
        {"page": "Policies",                 "views": 1},
        {"page": "FinOps NexGen Reports",    "views": 1},
    ]

# ─── PLATFORM METRICS (MongoDB — US East) ───────────────────────────────────

MONGO_CFG = {
    "host":        "52.154.142.32",
    "port":        27017,
    "username":    "demo",
    "password":    "Fd7Wv5ftLO5}k8",
    "auth_source": "admin",
}
MONGO_CONN_TIMEOUT_MS = 10000

# Try loading override from credentials.py
try:
    from credentials import MONGO_CFG as _MC  # type: ignore
    MONGO_CFG.update(_MC)
except (ImportError, AttributeError):
    pass

MONGO_DB = {
    "billing":   "billing_and_cost_analytics",
    "heatstack": "heatstack",
    "audit":     "audit_log",
}
MONGO_COLL = {
    "jobs":  "background_jobs",
    "sa":    "service_account_details",
    "audit": "request_audit",
}
MONGO_FIELD_MAP = {
    "background_job": {
        "payload_type":       "payload.__type",
        "payload_type_value": "BackgroundJobPayloadForCloudUsageBilling",
        "status":             "status",
        "status_completed":   "Completed",
        "updated_at":         "updated_at",
        "created_at":         "created_at",
        "service_account_id": "payload.service_account_id",
    },
    "service_account_details": {
        "id":        "_id",
        "tenant_id": "tenant_id",
    },
    "request_audit": {
        "created_at": "start_time",
        "duration":   "duration",
        "user":       "user_name",
        "slow_s":     30,
    },
}

AUDIT_HOURS     = 24
AUDIT_EXECUTOR  = "COST"
BILLING_PATHS   = [
    "/v1/billing_plans", "/v1/billing_plans/batch_definitions",
    "/v1/billing_plans/batch_versions", "/v1/cost/billing/aggregation/trend",
    "/v1/cost/estimated_cost", "/v1/internal/dimension/validate_grouping_rule_filters",
    "/v1/providers/billing/request_aggregate", "/v1/providers/billing/request_aggregate_trend",
    "/v1/providers/billing/request_rate_aggregate_trend",
    "/v1/providers/billing/request_usage_aggregate_trend",
    "/v2/billing/aggregation", "/v2/billing/aggregation/batch",
    "/v2/billing/aggregation/trend", "/v2/billing/extras",
    "/v2/billing/line_items_summary", "/v2/billing/platform/tags",
    "/v2/billing/tags", "/v2/billing_metrics/batch", "/v2/billing_metrics/list",
    "/v2/budget/dashboard/list_cloud_account_type", "/v2/budget/view/cloud_account",
    "/v2/budget/view/tenant", "/v2/budgets/threshold_alerts/view",
    "/v2/cost_anomaly/billing_cost_anomaly", "/v2/cost_anomaly/billing_cost_anomaly_resources",
    "/v3/budget/dashboard", "/v3/budget/dashboard/filters",
    "/v3/budget/dashboard/list_budgets", "/v3/budget/dashboard/list_currency",
    "/v3/budget/insights", "/v2/savings/filter", "/v2/savings/summary",
]
AUDIT_EXCLUDED_USERS = [
    "admin", "automation_in", "cs-metering", "automation_us", "automation_user",
    "automation_ingram", "validation", "qa_test", "apiautomation_produs",
    "validation2", "qa_user", "automation_mea",
]
AUDIT_INTERNAL_USERNAMES = {
    "admin.taylorfarms", "parthu_cs4cs", "cs-vidyasagar",
    "ganeshan-cs-qa", "admin.otsuka", "admin.convergetech", "blackstone",
}


def _mongo_get_field(doc: dict, dotpath: str):
    val = doc
    for part in dotpath.split("."):
        if not isinstance(val, dict):
            return None
        val = val.get(part)
    return val


def _make_mongo_client():
    try:
        from pymongo import MongoClient
    except ImportError:
        return None
    cfg = MONGO_CFG
    uri = (f"mongodb://{cfg['username']}:{cfg['password']}"
           f"@{cfg['host']}:{cfg['port']}"
           f"/?authSource={cfg['auth_source']}&directConnection=true")
    return MongoClient(uri, serverSelectionTimeoutMS=MONGO_CONN_TIMEOUT_MS,
                       socketTimeoutMS=30000)


def fetch_platform_metrics():
    """
    Fetches N-2 compliance and COST audit metrics from MongoDB (US East).
    Falls back to placeholder values if pymongo is not installed or connection fails.
    """
    import datetime as _dt

    now        = _dt.datetime.now(_dt.timezone.utc).replace(tzinfo=None)
    cutoff_24h = now - _dt.timedelta(hours=24)
    cutoff_48h = now - _dt.timedelta(hours=48)

    # ── N-2 job metrics ───────────────────────────────────────────────────────
    n2 = {
        "total_accounts": "—", "completed_jobs_24h": "—",
        "tenants_impacted": "—", "pending_n2_accts": "—",
        "older_backlog": "—", "compliance_pct": "—",
        "pending_names": [], "error": None,
    }
    # ── COST audit metrics ─────────────────────────────────────────────────────
    cost = {
        "total_requests_24h": "—", "max_response_sec": "—",
        "slow_requests_30s": "—", "users_impacted": "—",
        "external_users": "—", "internal_users": "—",
        "avg_slow_s": "—", "health_status": "—", "error": None,
    }

    client = None
    try:
        client = _make_mongo_client()
        if client is None:
            raise RuntimeError("pymongo not installed")
        # quick ping to catch VPN/network issues early
        client.admin.command("ping")
    except Exception as exc:
        err = f"MongoDB unavailable: {exc}"
        print(f"  [Mongo] {err}", file=sys.stderr)
        n2["error"]   = err
        cost["error"] = err
        if client:
            try: client.close()
            except: pass
        return {"n2": n2, "cost": cost}

    # ── Section 1: N-2 jobs ───────────────────────────────────────────────────
    try:
        jf = MONGO_FIELD_MAP["background_job"]
        sf = MONGO_FIELD_MAP["service_account_details"]
        n2_window = {jf["created_at"]: {"$gt": cutoff_48h, "$lte": cutoff_24h}}
        base      = {jf["payload_type"]: jf["payload_type_value"]}

        jobs = client[MONGO_DB["billing"]][MONGO_COLL["jobs"]]
        sa_c = client[MONGO_DB["heatstack"]][MONGO_COLL["sa"]]

        all_n2_ids  = jobs.distinct(jf["service_account_id"], {**base, **n2_window})
        total_accts = len([x for x in all_n2_ids if x is not None])

        completed_24h = jobs.count_documents({
            **base,
            jf["status"]:     jf["status_completed"],
            jf["updated_at"]: {"$gte": cutoff_24h},
        })

        n2_pending_docs = list(jobs.find(
            {**base, jf["status"]: {"$in": ["Ready", "Pending", "Waiting"]}, **n2_window},
            {jf["service_account_id"]: 1, "_id": 0},
        ))
        n2_pending_ids = {_mongo_get_field(d, jf["service_account_id"])
                          for d in n2_pending_docs
                          if _mongo_get_field(d, jf["service_account_id"]) is not None}

        older_backlog_docs = list(jobs.find(
            {**base, jf["status"]: {"$in": ["Ready", "Pending", "Waiting"]},
             jf["created_at"]: {"$lt": cutoff_48h}},
            {jf["service_account_id"]: 1, "_id": 0},
        ))
        older_backlog_ids = {_mongo_get_field(d, jf["service_account_id"])
                             for d in older_backlog_docs
                             if _mongo_get_field(d, jf["service_account_id"]) is not None}
        older_backlog = len(older_backlog_docs)

        combined_ids     = n2_pending_ids | older_backlog_ids
        tenants_impacted = 0
        pending_names    = []
        if combined_ids:
            sa_docs = list(sa_c.find(
                {sf["id"]: {"$in": list(combined_ids)}},
                {sf["id"]: 1, sf["tenant_id"]: 1, "name": 1},
            ))
            tenant_set       = {str(d[sf["tenant_id"]]) for d in sa_docs if d.get(sf["tenant_id"])}
            tenants_impacted = len(tenant_set)
            # names of N-2 pending accounts specifically
            n2_sa_docs = list(sa_c.find(
                {sf["id"]: {"$in": list(n2_pending_ids)}},
                {sf["id"]: 1, "name": 1},
            )) if n2_pending_ids else []
            pending_names = [d["name"] for d in n2_sa_docs if d.get("name")]

        n2_pending    = len(n2_pending_ids)
        compliance_pct = (0.0 if total_accts == 0 else
                          100.0 if n2_pending == 0 else
                          round((1 - n2_pending / total_accts) * 100, 1))

        n2.update({
            "total_accounts":    total_accts,
            "completed_jobs_24h": completed_24h,
            "tenants_impacted":  tenants_impacted,
            "pending_n2_accts":  n2_pending,
            "older_backlog":     older_backlog,
            "compliance_pct":    f"{compliance_pct}%",
            "pending_names":     pending_names,
        })
        print(f"  [Mongo] N-2: total={total_accts} pending={n2_pending} backlog={older_backlog} "
              f"compliance={compliance_pct}%")
    except Exception as exc:
        n2["error"] = str(exc)
        print(f"  [Mongo] N-2 error: {exc}", file=sys.stderr)

    # ── Section 2: COST audit metrics ─────────────────────────────────────────
    try:
        af     = MONGO_FIELD_MAP["request_audit"]
        slow_s = af["slow_s"]
        cutoff = now - _dt.timedelta(hours=AUDIT_HOURS)

        coll   = client[MONGO_DB["audit"]][MONGO_COLL["audit"]]
        time_q = {
            af["created_at"]: {"$gte": cutoff, "$lte": now},
            "executor":   AUDIT_EXECUTOR,
            "path":       {"$in": BILLING_PATHS},
            "user_name":  {"$exists": True, "$nin": AUDIT_EXCLUDED_USERS},
            "source_ip":  {"$ne": "127.0.0.1"},
        }

        u_field = f"${af['user']}"
        d_field = f"${af['duration']}"

        def _is_internal_expr(uf):
            return {"$or": [
                {"$regexMatch": {"input": uf, "regex": "^cs\\."}},
                {"$regexMatch": {"input": uf, "regex": "@corestack\\.io$"}},
                {"$in": [uf, list(AUDIT_INTERNAL_USERNAMES)]},
            ]}

        agg_result = list(coll.aggregate([
            {"$match": time_q},
            {"$group": {
                "_id":         None,
                "total":       {"$sum": 1},
                "max_dur":     {"$max": d_field},
                "slow_count":  {"$sum": {"$cond": [{"$gte": [d_field, slow_s]}, 1, 0]}},
                "slow_sum":    {"$sum": {"$cond": [{"$gte": [d_field, slow_s]}, d_field, 0]}},
                "ext_users":   {"$addToSet": {"$cond": {
                    "if":  {"$and": [{"$gte": [d_field, slow_s]},
                                     {"$not": [_is_internal_expr(u_field)]}]},
                    "then": u_field, "else": "$$REMOVE",
                }}},
                "int_users":   {"$addToSet": {"$cond": {
                    "if":  {"$and": [{"$gte": [d_field, slow_s]},
                                     _is_internal_expr(u_field)]},
                    "then": u_field, "else": "$$REMOVE",
                }}},
            }},
        ], allowDiskUse=True))

        if agg_result:
            row        = agg_result[0]
            slow_count = row.get("slow_count") or 0
            slow_sum   = row.get("slow_sum") or 0
            ext_list   = [u for u in (row.get("ext_users") or []) if u]
            int_list   = [u for u in (row.get("int_users") or []) if u]
            avg_slow   = round(slow_sum / slow_count, 2) if slow_count else 0.0
            max_s      = round(row.get("max_dur") or 0, 2)
            total_imp  = len(ext_list) + len(int_list)
            health     = "Bad" if slow_count > 0 else "Good"

            cost.update({
                "total_requests_24h": row.get("total") or 0,
                "max_response_sec":   max_s,
                "slow_requests_30s":  slow_count,
                "avg_slow_s":         avg_slow,
                "external_users":     len(ext_list),
                "internal_users":     len(int_list),
                "users_impacted":     total_imp,
                "health_status":      health,
            })
            print(f"  [Mongo] COST: requests={cost['total_requests_24h']} "
                  f"slow={slow_count} max={max_s}s health={health}")
        else:
            cost.update({
                "total_requests_24h": 0, "max_response_sec": 0.0,
                "slow_requests_30s": 0, "avg_slow_s": 0.0,
                "external_users": 0, "internal_users": 0,
                "users_impacted": 0, "health_status": "Good",
            })
    except Exception as exc:
        cost["error"] = str(exc)
        print(f"  [Mongo] COST error: {exc}", file=sys.stderr)
    finally:
        try: client.close()
        except: pass

    return {"n2": n2, "cost": cost}

# ─── CHART GENERATORS ────────────────────────────────────────────────────────

BRAND_COLORS = ["#1F6FBF", "#2EA8CC", "#5CC8A0", "#F4A460", "#E05C5C", "#9B59B6", "#F39C12", "#1ABC9C"]
FONT_FAMILY  = "DejaVu Sans"

def _chart_png(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def chart_pendo_top_pages(pages: list) -> io.BytesIO:
    """Horizontal bar chart — top pages by views."""
    names  = [p["page"] for p in pages]
    views  = [p["views"] for p in pages]
    colors = [BRAND_COLORS[i % len(BRAND_COLORS)] for i in range(len(names))]

    fig, ax = plt.subplots(figsize=(6, max(2.5, len(names) * 0.45)))
    bars = ax.barh(names[::-1], views[::-1], color=colors[::-1], edgecolor="none", height=0.6)
    for bar, val in zip(bars, views[::-1]):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height() / 2,
                str(val), va="center", ha="left", fontsize=9, color="#444")
    ax.set_xlabel("Page views", fontsize=9)
    ax.set_title("Top pages visited", fontsize=11, fontweight="bold", pad=10, color="#1F2937")
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.tick_params(axis="y", labelsize=9)
    ax.tick_params(axis="x", labelsize=8)
    ax.set_xlim(0, max(views) * 1.25 if views else 10)
    ax.xaxis.grid(True, linestyle="--", alpha=0.4)
    ax.set_axisbelow(True)
    fig.tight_layout()
    return _chart_png(fig)


def chart_pendo_visitor_activity(visitors: list) -> io.BytesIO:
    """Bar chart — events per active visitor (only those with events)."""
    active = [(v["visitor"].title(), v["events"]) for v in visitors if isinstance(v["events"], int)]
    if not active:
        # Return a simple "no activity" placeholder chart
        fig, ax = plt.subplots(figsize=(5, 1.5))
        ax.text(0.5, 0.5, "No visitor activity in this window",
                ha="center", va="center", fontsize=10, color="#999", transform=ax.transAxes)
        ax.axis("off")
        return _chart_png(fig)

    names  = [a[0] for a in active]
    events = [a[1] for a in active]

    fig, ax = plt.subplots(figsize=(max(4, len(names) * 1.2), 3.2))
    bars = ax.bar(names, events, color=BRAND_COLORS[0], edgecolor="none", width=0.55)
    for bar, val in zip(bars, events):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                str(val), ha="center", va="bottom", fontsize=9, color="#444")
    ax.set_ylabel("Events (3-day window)", fontsize=9)
    ax.set_title("Visitor activity", fontsize=11, fontweight="bold", pad=10, color="#1F2937")
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(axis="x", labelsize=9, rotation=15)
    ax.tick_params(axis="y", labelsize=8)
    ax.yaxis.grid(True, linestyle="--", alpha=0.4)
    ax.set_axisbelow(True)
    fig.tight_layout()
    return _chart_png(fig)


def chart_ado_by_state(ado_items: list) -> io.BytesIO:
    """Donut chart — work items by state."""
    state_order  = ["New", "In Progress", "Awaiting Deployment"]
    state_colors = {"New": "#2EA8CC", "In Progress": "#1F6FBF", "Awaiting Deployment": "#5CC8A0"}

    counts = {s: 0 for s in state_order}
    for item in ado_items:
        s = item["state"]
        if s in counts:
            counts[s] += 1

    labels = [s for s in state_order if counts[s] > 0]
    sizes  = [counts[s] for s in labels]
    colors = [state_colors[s] for s in labels]

    fig, ax = plt.subplots(figsize=(4.5, 3.2))
    if not sizes:
        ax.text(0.5, 0.5, "No items", ha="center", va="center",
                fontsize=11, color="#999", transform=ax.transAxes)
        ax.axis("off")
    else:
        wedges, texts, autotexts = ax.pie(
            sizes, labels=None, colors=colors, autopct="%1.0f%%",
            startangle=90, wedgeprops={"width": 0.55, "edgecolor": "white", "linewidth": 2},
            pctdistance=0.75,
        )
        for at in autotexts:
            at.set_fontsize(10)
            at.set_color("white")
            at.set_fontweight("bold")
        legend_patches = [mpatches.Patch(color=colors[i], label=f"{labels[i]} ({sizes[i]})")
                          for i in range(len(labels))]
        ax.legend(handles=legend_patches, loc="lower center", bbox_to_anchor=(0.5, -0.18),
                  ncol=len(labels), fontsize=8.5, frameon=False)
    ax.set_title("ADO items by state", fontsize=11, fontweight="bold", pad=10, color="#1F2937")
    fig.tight_layout()
    return _chart_png(fig)


def chart_ado_by_priority(ado_items: list) -> io.BytesIO:
    """Bar chart — work items by priority."""
    priority_map = {1: "P1 – Critical", 2: "P2 – High", 3: "P3 – Medium", 4: "P4 – Low"}
    pri_colors   = {"P1 – Critical": "#E05C5C", "P2 – High": "#F4A460",
                    "P3 – Medium": "#2EA8CC",   "P4 – Low": "#5CC8A0"}

    counts: dict = {}
    for item in ado_items:
        label = priority_map.get(item["priority"], f"P{item['priority']}" if item["priority"] else "—")
        counts[label] = counts.get(label, 0) + 1

    ordered = [p for p in priority_map.values() if p in counts]
    values  = [counts[p] for p in ordered]
    colors  = [pri_colors.get(p, "#aaa") for p in ordered]

    fig, ax = plt.subplots(figsize=(4.5, 3.2))
    if not ordered:
        ax.text(0.5, 0.5, "No items", ha="center", va="center",
                fontsize=11, color="#999", transform=ax.transAxes)
        ax.axis("off")
    else:
        bars = ax.bar(ordered, values, color=colors, edgecolor="none", width=0.5)
        for bar, val in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05,
                    str(val), ha="center", va="bottom", fontsize=10, fontweight="bold", color="#444")
        ax.set_ylabel("Work items", fontsize=9)
        ax.spines[["top", "right"]].set_visible(False)
        ax.tick_params(axis="x", labelsize=8.5)
        ax.tick_params(axis="y", labelsize=8)
        ax.yaxis.grid(True, linestyle="--", alpha=0.4)
        ax.set_axisbelow(True)
        ax.set_ylim(0, max(values) + 1)
    ax.set_title("ADO items by priority", fontsize=11, fontweight="bold", pad=10, color="#1F2937")
    fig.tight_layout()
    return _chart_png(fig)


# ─── DOCX HELPERS ────────────────────────────────────────────────────────────

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand colours
C_DARK   = (0x1F, 0x29, 0x37)
C_WHITE  = (0xFF, 0xFF, 0xFF)
C_ACCENT = (0x00, 0x70, 0xC0)
C_GREY   = (0xF2, 0xF2, 0xF2)
C_MUTED  = (0x75, 0x75, 0x75)
C_GREEN  = (0x70, 0xAD, 0x47)
C_RED    = (0xFF, 0x00, 0x00)


def _set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, border_side="bottom", color="CCCCCC", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = tcPr.find(qn("w:tcBdr"))
    if tcBdr is None:
        tcBdr = OxmlElement("w:tcBdr")
        tcPr.append(tcBdr)
    side = OxmlElement(f"w:{border_side}")
    side.set(qn("w:val"),   "single")
    side.set(qn("w:sz"),    sz)
    side.set(qn("w:space"), "0")
    side.set(qn("w:color"), color)
    tcBdr.append(side)


def _para_fmt(para, space_before=0, space_after=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.alignment = align


def _run(para, text, bold=False, italic=False, size=10,
         color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _heading(doc, text, level_size=12, color=C_DARK):
    p = doc.add_paragraph()
    _para_fmt(p, space_before=10, space_after=2)
    _run(p, text, bold=True, size=level_size, color=color)
    return p


def _add_table(doc, headers, rows, col_widths=None, zebra=True, header_bg=C_DARK):
    """Generic styled table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        _run(p, h, bold=True, size=9, color=C_WHITE)
        _set_cell_bg(hdr_cells[i], header_bg)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        bg = C_GREY if (zebra and r_idx % 2 == 0) else C_WHITE
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            _run(p, str(val), size=9)
            _set_cell_bg(cells[c_idx], bg)

    # Column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return tbl


# ─── REPORT RENDERER ─────────────────────────────────────────────────────────

def render_docx(ado_items, eu_visitors, useast_visitors, other_visitors, pages, metrics):
    """
    Single-page scrum report matching the June 17 layout.
    Pendo section shows all visitors (EU + USEast combined) with a Region column
    and an EU / USEast count summary. ADO incidents appear once.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    today    = datetime.date.today()
    now_ist  = datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(hours=5, minutes=30)
    date_str = fmt_date(today)
    time_str = now_ist.strftime("%I:%M %p IST").lstrip("0")

    pendo_end   = today
    pendo_start = today - datetime.timedelta(days=PENDO_WINDOW_DAYS - 1)
    pendo_range = f"{fmt_mon(pendo_start)}–{fmt_mon(pendo_end)}"

    n2  = metrics["n2"]
    svc = metrics["cost"]
    open_incidents = [i for i in ado_items if i["state"] in ("New", "In Progress", "Awaiting Deployment")]

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    _para_fmt(title_p, space_before=0, space_after=2)
    _run(title_p, "Blackstone — daily scrum update", bold=True, size=16, color=C_DARK)

    sub_p = doc.add_paragraph()
    _para_fmt(sub_p, space_before=0, space_after=6)
    _run(sub_p, f"{date_str}   ·   Report generated {time_str}", size=9, color=C_MUTED)

    # ── At a glance KPI ───────────────────────────────────────────────────────
    _heading(doc, "At a glance")
    kpi_tbl = doc.add_table(rows=2, cols=4)
    kpi_tbl.style = "Table Grid"
    kpi_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    kpi_headers = ["Compliance", "API requests", "Slow req ≥30s", "Open incidents"]
    open_ids    = ", ".join([f"#{i['id']}" for i in open_incidents[:3]]) or "None"
    kpi_vals    = [
        n2.get("compliance_pct", "—"),
        f"{svc.get('total_requests_24h','—')}\nMax {svc.get('max_response_sec','—')}s resp.",
        str(svc.get("slow_requests_30s", "—")),
        f"{len(open_incidents)}\n{open_ids}",
    ]
    for i, h in enumerate(kpi_headers):
        c = kpi_tbl.rows[0].cells[i]
        c.text = ""
        _run(c.paragraphs[0], h, bold=True, size=9, color=C_WHITE)
        _set_cell_bg(c, C_DARK)
    for i, v in enumerate(kpi_vals):
        c = kpi_tbl.rows[1].cells[i]
        c.text = ""
        _run(c.paragraphs[0], v, bold=True, size=13, color=C_ACCENT)
        _set_cell_bg(c, C_GREY)
    kpi_w = [1.5, 1.8, 1.8, 1.8]
    for row in kpi_tbl.rows:
        for i, w in enumerate(kpi_w):
            row.cells[i].width = Inches(w)
    doc.add_paragraph()

    # ── Platform performance ──────────────────────────────────────────────────
    _heading(doc, "Platform performance  (N-2 metric — US East)")
    n2_subtitle = doc.add_paragraph()
    _para_fmt(n2_subtitle, space_before=0, space_after=4)
    _run(n2_subtitle,
         "N-2 window = billing jobs created 24h–48h ago  ·  "
         "Compliance % = (1 – Pending N-2 Accts ÷ Total Accts) × 100",
         size=8, italic=True, color=C_MUTED)

    if n2.get("error"):
        err_p = doc.add_paragraph()
        _para_fmt(err_p, space_before=0, space_after=4)
        _run(err_p, f"⚠ MongoDB unavailable — {n2['error'][:120]}", size=9, color=C_RED)
    else:
        pending_names = n2.get("pending_names", [])
        pending_val   = str(n2["pending_n2_accts"])
        if pending_names:
            pending_val += "  (" + ", ".join(pending_names) + ")"
        _add_table(doc,
            headers=["Metric", "Value"],
            rows=[
                ("Total accounts (N-2 window)",  n2["total_accounts"]),
                ("Completed jobs (24h)",          n2["completed_jobs_24h"]),
                ("Tenants impacted",              n2["tenants_impacted"]),
                ("Pending N-2 accounts",          pending_val),
                ("Older backlog (>48h queued)",   n2["older_backlog"]),
                ("Compliance %",                  n2["compliance_pct"]),
            ],
            col_widths=[3.5, 2.8],
        )
    doc.add_paragraph()

    # ── Service metric (COST) ─────────────────────────────────────────────────
    _heading(doc, "Service metric (COST — US East)")
    svc_subtitle = doc.add_paragraph()
    _para_fmt(svc_subtitle, space_before=0, space_after=4)
    _run(svc_subtitle,
         f"Executor: COST  ·  Cost/billing endpoints only  ·  "
         f"Real users only (system & automation excluded)  ·  Slow threshold ≥30s  ·  Last {AUDIT_HOURS}h",
         size=8, italic=True, color=C_MUTED)

    if svc.get("error"):
        err_p = doc.add_paragraph()
        _para_fmt(err_p, space_before=0, space_after=4)
        _run(err_p, f"⚠ MongoDB unavailable — {svc['error'][:120]}", size=9, color=C_RED)
    else:
        avg_slow_str = (f"{svc['avg_slow_s']} sec" if svc.get("slow_requests_30s", 0) else "—")
        _add_table(doc,
            headers=["Metric", "Value"],
            rows=[
                (f"Total requests (last {AUDIT_HOURS}h)",  svc["total_requests_24h"]),
                ("Max response time",                       f"{svc['max_response_sec']} sec"),
                ("Avg slow request duration",               avg_slow_str),
                ("Slow requests ≥30s",                      svc["slow_requests_30s"]),
                ("External users impacted (slow req)",      svc["external_users"]),
                ("Internal users impacted (slow req)",      svc["internal_users"]),
                ("Health status",                           svc["health_status"]),
            ],
            col_widths=[3.5, 2.8],
        )
    doc.add_paragraph()

    # ── Pendo engagement (all visitors; EU + USEast counted separately) ──────
    all_visitors = eu_visitors + useast_visitors + other_visitors
    _heading(doc, f"Pendo engagement — Blackstone segment  ·  last {PENDO_WINDOW_DAYS} days ({pendo_range})")

    note_p = doc.add_paragraph()
    _para_fmt(note_p, space_before=0, space_after=4)
    _run(note_p,
         f"Source: Blackstone Pendo segment only.  {PENDO_EXCLUDED_EMAIL} excluded from all counts.  "
         f"Grey rows = segment members with no activity in this window.",
         size=8, italic=True, color=C_MUTED)

    # EU / USEast count summary line
    trend_p = doc.add_paragraph()
    _para_fmt(trend_p, space_before=0, space_after=6)
    _run(trend_p, f"EU (portal.corestack.io): ", bold=True, size=9, color=C_DARK)
    _run(trend_p, f"{len(eu_visitors)} active visitor{'s' if len(eu_visitors) != 1 else ''}",
         size=9, color=C_ACCENT)
    _run(trend_p, "     USEast (useast.corestack.io): ", bold=True, size=9, color=C_DARK)
    _run(trend_p, f"{len(useast_visitors)} active visitor{'s' if len(useast_visitors) != 1 else ''}",
         size=9, color=C_ACCENT)

    _heading(doc, "Active visitors", level_size=10)
    visitor_rows = []
    for v in all_visitors:
        ev = str(v["events"])     if isinstance(v["events"],     int) else "-"
        da = str(v["daysActive"]) if isinstance(v["daysActive"], int) else "-"
        mn = str(v["minutes"])    if isinstance(v["minutes"],    int) else "-"
        visitor_rows.append((v["visitor"], v["domain"], v.get("account", ""), ev, da, mn, v["lastSeen"]))

    _add_table(doc,
        headers=["Visitor (Blackstone segment)", "Domain", "Account", "Events (3d)", "Days active", "Minutes", "Last seen"],
        rows=visitor_rows or [("—", "No activity in this window.", "", "", "", "", "")],
        col_widths=[1.7, 1.5, 1.2, 0.9, 0.9, 0.7, 0.9],
    )
    doc.add_paragraph()

    activity_img = chart_pendo_visitor_activity(all_visitors)
    pages_img    = chart_pendo_top_pages(pages)

    chart_tbl = doc.add_table(rows=1, cols=2)
    chart_tbl.style = "Table Grid"
    chart_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    left_cell  = chart_tbl.rows[0].cells[0]
    right_cell = chart_tbl.rows[0].cells[1]
    _set_cell_bg(left_cell,  C_WHITE)
    _set_cell_bg(right_cell, C_WHITE)
    left_cell.paragraphs[0].add_run().add_picture(activity_img, width=Inches(3.3))
    right_cell.paragraphs[0].add_run().add_picture(pages_img,   width=Inches(3.3))
    for cell in [left_cell, right_cell]:
        cell.width = Inches(3.4)
    doc.add_paragraph()

    _heading(doc, "Top pages visited", level_size=10)
    _add_table(doc,
        headers=["Page", "Views"],
        rows=[(p["page"], p["views"]) for p in pages] or [("—", "—")],
        col_widths=[4.5, 0.8],
    )
    doc.add_paragraph()

    # ── ADO incidents ─────────────────────────────────────────────────────────
    ado_url = (
        f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}/_workitems?"
        f"filter=tags+eq+'Blackstone'"
    )
    _heading(doc, "Blackstone incidents — ADO status")
    link_p = doc.add_paragraph()
    _para_fmt(link_p, space_before=0, space_after=4)
    _run(link_p, "ADO query: ", size=9, color=C_MUTED)
    _run(link_p, ado_url, size=9, color=C_ACCENT)

    ado_rows = []
    for item in ado_items:
        pri = f"P{item['priority']}" if item["priority"] else "—"
        ado_rows.append((
            str(item["id"]),
            item["title"],
            item["assignedTo"],
            item.get("areaPath", "").split("\\")[-1] or "—",
            pri,
            item["state"],
            item["workItemType"],
        ))
    if not ado_rows:
        ado_rows = [("—", "No active Blackstone incidents found.", "", "", "", "", "")]

    _add_table(doc,
        headers=["ID", "Title", "Assigned to", "Bundle", "Pri", "State", "Classification"],
        rows=ado_rows,
        col_widths=[0.6, 2.8, 1.4, 0.9, 0.4, 1.0, 0.7],
    )
    doc.add_paragraph()

    state_img = chart_ado_by_state(ado_items)
    pri_img   = chart_ado_by_priority(ado_items)

    ado_chart_tbl = doc.add_table(rows=1, cols=2)
    ado_chart_tbl.style = "Table Grid"
    ado_chart_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    sc = ado_chart_tbl.rows[0].cells[0]
    pc = ado_chart_tbl.rows[0].cells[1]
    _set_cell_bg(sc, C_WHITE)
    _set_cell_bg(pc, C_WHITE)
    sc.paragraphs[0].add_run().add_picture(state_img, width=Inches(3.3))
    pc.paragraphs[0].add_run().add_picture(pri_img,   width=Inches(3.3))
    for cell in [sc, pc]:
        cell.width = Inches(3.4)

    # ── Footer ────────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    _para_fmt(footer_p, space_before=12, space_after=0)
    _run(footer_p,
         "CoreStack · Cloud Operations Intelligence  ·  Daily scrum report  ·  "
         "N-2 = jobs created 24h–48h ago  ·  Slow = request ≥30s",
         size=8, color=C_MUTED, italic=True)

    return doc


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    using_fake = []
    if not ADO_PAT:
        using_fake.append("ADO   → using FAKE data  (add ADO_PAT to credentials.py)")
    else:
        print("✓ ADO PAT detected — will fetch live work items")

    if not PENDO_API_KEY:
        using_fake.append("Pendo → using FAKE data  (set env var: export PENDO_API_KEY=...)")
    else:
        print("✓ Pendo API key detected — will fetch live visitors")

    if using_fake:
        print("\n⚠️  WARNING — missing credentials:")
        for msg in using_fake:
            print(f"   {msg}")
        print()

    print("Fetching ADO Blackstone work items...")
    ado_items = fetch_ado_blackstone_incidents()

    print("Fetching Pendo engagement...")
    accounts     = _fetch_blackstone_accounts()
    all_visitors = fetch_pendo_all_visitors(accounts=accounts)
    eu_visitors, useast_visitors, other_visitors = split_visitors_by_region(all_visitors)
    print(f"  → EU visitors: {len(eu_visitors)}  |  USEast visitors: {len(useast_visitors)}  |  unclassified: {len(other_visitors)}")

    print("Fetching platform metrics...")
    metrics = fetch_platform_metrics()

    pages = fetch_pendo_top_pages(accounts=accounts)

    doc = render_docx(ado_items, eu_visitors, useast_visitors, other_visitors, pages, metrics)

    out_file = f"Blackstone_Scrum_{datetime.date.today().strftime('%Y%m%d')}.docx"
    doc.save(out_file)
    print(f"\n✓ Saved to {out_file}")


if __name__ == "__main__":
    main()

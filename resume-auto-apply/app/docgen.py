from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.models import Job, Profile


def generate_resume_docx(profile: Profile, tailored: dict, out_path: Path) -> None:
    doc = Document()

    title = doc.add_heading(profile.full_name or "Resume", level=0)
    title.alignment = 1

    contact_bits = [b for b in [profile.email, profile.phone] if b]
    if contact_bits:
        p = doc.add_paragraph(" | ".join(contact_bits))
        p.alignment = 1

    if tailored.get("headline"):
        p = doc.add_paragraph()
        run = p.add_run(tailored["headline"])
        run.bold = True
        run.font.size = Pt(13)
        p.alignment = 1

    if tailored.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(tailored["summary"])

    if tailored.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(" | ".join(tailored["skills"]))

    if tailored.get("experience"):
        doc.add_heading("Experience", level=1)
        for role in tailored["experience"]:
            header = f"{role.get('title', '')} — {role.get('company', '')}".strip(" —")
            p = doc.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            if role.get("dates"):
                p.add_run(f"  ({role['dates']})")
            for bullet in role.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    if tailored.get("education"):
        doc.add_heading("Education", level=1)
        for edu in tailored["education"]:
            doc.add_paragraph(edu, style="List Bullet")

    doc.save(out_path)


def generate_cover_letter_docx(profile: Profile, job: Job, cover_letter_text: str, out_path: Path) -> None:
    doc = Document()
    if profile.full_name:
        doc.add_paragraph(profile.full_name)
    contact_bits = [b for b in [profile.email, profile.phone] if b]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))
    doc.add_paragraph("")
    doc.add_paragraph(f"Re: {job.title} at {job.company}")
    doc.add_paragraph("")
    for paragraph in cover_letter_text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.save(out_path)
